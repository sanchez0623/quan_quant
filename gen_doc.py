# -*- coding: utf-8 -*-
"""生成《A股个人量化回测系统 完整实现方案》docx 文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
GRAY = RGBColor(0x59, 0x59, 0x59)

doc = Document()

# ---------- 全局样式 ----------
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for i, (size, color) in enumerate([(20, DARK_BLUE), (15, DARK_BLUE), (12.5, DARK_BLUE), (11, DARK_BLUE)], start=1):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Calibri'
    h.font.size = Pt(size)
    h.font.color.rgb = color
    h.font.bold = True
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def p(text, bold=False, size=10.5, color=None, italic=False, style_name=None):
    para = doc.add_paragraph(style=style_name)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return para

def bullet(text, level=0):
    para = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    para.add_run(text)
    return para

def numbered(text):
    para = doc.add_paragraph(style='List Number')
    para.add_run(text)
    return para

def code(text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)
    return para

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(9.5)
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = str(v)
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(9)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    doc.add_paragraph()
    return t

# ---------- 封面 ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('\nA股个人量化回测系统\n完整实现方案')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = DARK_BLUE
run.font.name = 'Calibri'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run(f'版本 V1.0    日期 {datetime.date.today().strftime("%Y-%m-%d")}')
run.font.size = Pt(12)
run.font.color.rgb = GRAY

doc.add_paragraph()
p('定位：面向个人及少量朋友使用的 A 股量化回测平台（不接真实券商），支持日线与 5 分钟线策略回测、加减仓与做 T、参数风控、完善统计，并接入 AI 大模型进行评估与调优。', color=GRAY)
doc.add_page_break()

# ---------- 目录说明 ----------
doc.add_heading('文档目录', level=1)
toc = [
    '1. 项目概述', '2. 总体架构', '3. 技术选型', '4. 数据层设计',
    '5. 回测引擎设计', '6. AI 辅助调优（多 LLM 配置）', '7. 前后端设计',
    '8. 部署方案（阿里云 + Docker + 外网访问）', '9. 安全与配置管理（GitHub 开源安全）',
    '10. 性能优化', '11. 断点续测策略', '12. 实施路线图', '13. 关键设计决策汇总',
]
for t in toc:
    p('    ' + t)

doc.add_page_break()

# ================= 1. 项目概述 =================
doc.add_heading('1. 项目概述', level=1)
doc.add_heading('1.1 项目目标', level=2)
bullet('构建一个 A 股量化回测系统，仅用于策略回测与研究，不接真实券商、不做实盘下单。')
bullet('支持日线与 5 分钟线两个周期的策略回测。')
bullet('完整支持加减仓、做 T（底仓当日回转）、分层持仓、参数化风控。')
bullet('输出完善的绩效统计与全过程交易记录，并可接入 AI 大模型进行评估与策略参数优化。')
bullet('部署于阿里云，本人与几位异地的朋友通过公网访问使用。')

doc.add_heading('1.2 使用规模与约束', level=2)
table(
    ['项目', '说明'],
    [
        ['用户规模', '个人 + 最多几位朋友，并发回测任务 ≤ 5'],
        ['数据范围', '全 A 约 5400 只股票，历史 5 年（一次性全量拉取 + 每日增量更新）'],
        ['K 线周期', '日线 + 5 分钟线（不需要 1 分钟线）'],
        ['部署环境', '阿里云轻量应用服务器（国内节点，无代理环境）'],
        ['开发环境', '本地 Windows + Clash Verge 代理（需特殊处理，见 4.6 节）'],
        ['代码托管', 'GitHub（开源或私有仓库），严禁泄露 API Key 等敏感信息'],
    ],
    widths=[4, 12],
)

doc.add_heading('1.3 不做的事情（明确边界）', level=2)
bullet('不做实盘交易、不接券商 API。')
bullet('不做 1 分钟线及 Tick 级数据。')
bullet('不做消息队列（Celery/Redis）、K8s 等重型基础设施——个人规模用不上。')
bullet('不做精细到 bar 级的单次回测断点续跑（理由与替代方案见第 11 章）。')

doc.add_page_break()

# ================= 2. 总体架构 =================
doc.add_heading('2. 总体架构', level=1)
p('系统采用前后端分离架构，整体分五层：')
code(
'┌───────────────────── 前端 (React + TypeScript) ─────────────────────┐\n'
'│  策略管理 │ 回测控制台 │ K线+交易标记 │ 统计报告 │ 寻优面板 │ AI面板 │\n'
'└────────────────────────────┬────────────────────────────────────┘\n'
'                     REST API + WebSocket(进度推送)\n'
'┌────────────────────────────┴────────────────────────────────────┐\n'
'│                     后端 API 层 (FastAPI)                         │\n'
'│   策略管理 │ 回测任务 │ 数据查询 │ 报告统计 │ LLM调度 │ 用户认证   │\n'
'└──────┬──────────────┬────────────────────┬─────────────────────┘\n'
'       │              │                    │\n'
'┌──────┴──────┐ ┌─────┴──────┐      ┌──────┴───────┐\n'
'│ 任务执行层   │ │ 回测引擎    │      │  AI 分析层    │\n'
'│ 进程池+队列  │ │ 事件驱动撮合 │      │ 多LLM+Optuna │\n'
'└──────┬──────┘ └─────┬──────┘      └──────┬───────┘\n'
'       │              │                    │\n'
'┌──────┴──────────────┴────────────────────┴─────────────────────┐\n'
'│ 数据层: 数据更新服务(多源自动切换) → Parquet数据湖 + DuckDB/SQLite │\n'
'│       (日线/5分钟线/复权因子/交易日历/股票基础信息)                │\n'
'└────────────────────────────────────────────────────────────────┘'
)
doc.add_heading('2.1 架构分层说明', level=2)
table(
    ['层次', '组件', '职责'],
    [
        ['前端', 'React + AntD + KLineCharts', '策略编辑、回测发起、结果可视化、AI 交互'],
        ['API 层', 'FastAPI', '统一 REST 接口、JWT 认证、WebSocket 进度推送、静态文件托管'],
        ['任务执行层', '进程池 + 内存任务队列', '回测/寻优任务异步执行，避免阻塞 API'],
        ['回测引擎', '自研事件驱动引擎', '撮合、账户、订单、风控、T+1/涨跌停模拟、统计输出'],
        ['AI 分析层', '多 LLM Provider + Optuna', '报告解读、参数搜索空间建议、贝叶斯寻优'],
        ['数据层', 'Parquet + DuckDB + 数据更新服务', '行情存储、查询、多数据源增量更新与自动切换'],
    ],
    widths=[3, 5, 8],
)
p('个人规模简化决策：不引入 PostgreSQL / Celery / Redis，元数据用 SQLite（或 DuckDB 建表），任务队列用 FastAPI BackgroundTasks + multiprocessing 进程池实现，缓存用进程内字典 / diskcache。', color=GRAY)

doc.add_page_break()

# ================= 3. 技术选型 =================
doc.add_heading('3. 技术选型', level=1)
doc.add_heading('3.1 回测引擎：自研事件驱动 + 向量化信号', level=2)
p('对主流框架的评估结论：')
table(
    ['方案', '优点', '缺点', '结论'],
    [
        ['backtrader', '事件驱动、订单模型完善、中文资料多', '性能一般、维护停滞、做T/分层持仓需大量 hack'],
        ['vnpy', '国内实盘生态好', '偏实盘，回测统计弱，过重'],
        ['qlib (微软)', 'AI/因子研究一流', '偏因子挖掘，交易过程模拟弱'],
        ['RQAlpha', 'A股规则适配好', '维护放缓，定制成本高'],
        ['自研轻量引擎', '完全贴合做T/加减仓/动态风控需求，可产品化，可嵌入AI闭环，可输出结构化交易过程', '有约2000~3000行核心开发量'],
    ],
    widths=[3, 5, 6, 2.5],
)
p('推荐自研，核心理由：')
numbered('做 T、分层加减仓、T+1、动态阈值风控是本系统核心需求，第三方框架均需大量魔改；')
numbered('引擎只需做"撮合 + 账户 + 订单"三件事，信号计算交给 polars 向量化，工作量可控；')
numbered('AI 闭环需要引擎输出结构化交易过程数据（每笔交易的理由、归属、贡献分解），自研引擎可直接埋点。')

doc.add_heading('3.2 技术栈清单', level=2)
table(
    ['类别', '选型', '说明'],
    [
        ['后端框架', 'FastAPI + Uvicorn', '异步、自动文档、WebSocket 原生支持'],
        ['前端', 'React + TypeScript + Ant Design', 'K线用 KLineCharts（交易标记支持好）'],
        ['数据处理', 'polars（主）+ pandas（兼容）', 'polars 多线程，指标计算比 pandas 快 5~10 倍'],
        ['行情存储', 'Parquet 文件 + DuckDB 查询', '3.2亿行级 5 分钟线查询秒回，零运维'],
        ['元数据', 'SQLite', '策略、任务、报告、用户，个人规模够用'],
        ['任务执行', 'BackgroundTasks + multiprocessing 进程池', '回测长任务异步化'],
        ['参数寻优', 'Optuna', '贝叶斯优化、自带剪枝、SQLite 存储天然支持断点续跑'],
        ['AI 接口', 'OpenAI 兼容协议多 Provider 抽象', '硅基流动/DeepSeek/智谱/Ollama 等，见第 6 章'],
        ['反向代理', 'Caddy', '自动 HTTPS 证书，零配置'],
        ['容器化', 'Docker + docker-compose', '单 compose 文件，本地与云端一致'],
        ['加速（可选）', 'Numba JIT', '撮合核心循环编译，分钟线大股票池场景使用'],
    ],
    widths=[3.2, 5.3, 7.5],
)

doc.add_page_break()

# ================= 4. 数据层设计 =================
doc.add_heading('4. 数据层设计', level=1)
doc.add_heading('4.1 数据范围与规模', level=2)
table(
    ['数据类型', '规模估算', '存储量'],
    [
        ['日线（全A 5400只 × 5年）', '约 660 万行', '单 Parquet 文件约 200MB'],
        ['5分钟线（全A × 5年）', '约 3.2 亿行', '按股票分文件，约 40~50GB'],
        ['复权因子', '日线级', '几十 MB'],
        ['交易日历 / 股票列表 / ST状态', '元数据级', '极小'],
    ],
    widths=[6, 4, 6],
)

doc.add_heading('4.2 存储结构', level=2)
code(
'/data\n'
'  daily.parquet            # 全部日线（一个文件）\n'
'  minute5/{code}.parquet   # 5分钟线按股票分文件（每只约10MB）\n'
'  adj_factor.parquet       # 复权因子\n'
'  trade_calendar.parquet   # 交易日历\n'
'  stock_basic.parquet      # 股票基础信息(含ST状态)\n'
'  meta.db                  # SQLite: 更新水位/任务/报告等元数据'
)
p('要点：')
bullet('存储用不复权原始价 + 复权因子，回测时按需计算。')
bullet('回测内部统一使用后复权价——前复权会因除权除息导致历史数据漂移，回测结果不可复现；报告展示时再换算回真实价格。')
bullet('回测引擎按需 lazy-load（只读策略池涉及的股票），polars 读 Parquet。')

doc.add_heading('4.3 数据源体系（多源 + 自动切换）', level=2)
p('数据源评估与分工：')
table(
    ['数据源', '5年5分钟线', '评价', '定位'],
    [
        ['mootdx（通达信协议）', '有', '直连通达信行情服务器，速度快、稳定性好、分钟线历史全，不受网页改版影响。注意：5分钟历史数据在"扩展行情服务器"，配置时需指定', '主源：分钟线'],
        ['baostock', '有', '稳定、数据干净，但慢且官方要求基本串行；日线量小不怕慢', '主源：日线/复权因子'],
        ['东方财富（efinance / akshare em接口）', '部分', '快，但本质爬虫，接口随时可能因网站改版失效', '备源 + 校验基准'],
        ['AKShare', '部分', '聚合爬虫源，覆盖广但波动大', '备源；股票列表/ST状态/涨跌停价'],
        ['腾讯财经', '无历史', '偏实时行情', '仅补当日数据（可选）'],
        ['理杏仁', '受限', '强在基本面/估值，K线历史有额度限制且收费', '不做行情源'],
    ],
    widths=[4.5, 2, 6.5, 3],
)
p('推荐组合：')
code(
'分钟线:   mootdx(主) → baostock(备)\n'
'日线:     baostock(主) → akshare/东财(备)\n'
'复权因子: baostock(主) → 东财(备，校验用)\n'
'基础信息: akshare(股票列表/ST状态/涨跌停价)')

doc.add_heading('4.4 数据源抽象与自动切换', level=2)
p('设计统一的 DataSource 抽象接口，各数据源实现同一套方法，更新服务内置调度逻辑：')
code(
'DataSource (抽象接口)\n'
'  ├── MootdxSource      # get_daily / get_minute5 / get_adj_factor ...\n'
'  ├── BaostockSource\n'
'  ├── AkshareSource\n'
'  └── EastmoneySource')
p('调度策略：')
numbered('健康检查：每次更新前发轻量探针请求（如拉一日日线），超时 10 秒判故障。')
numbered('优先级降级：主源失败 → 自动切备源 → 记录日志 + 通知（可接微信推送/邮件）。')
numbered('按数据类型分别指定主备源（日线与分钟线走不同链路）。')
numbered('数据对齐校验：切换数据源后，重叠区间抽样比对价格，防止口径不一致（复权因子差异等）。')
numbered('恢复检测：后台每 30 分钟探测主源，恢复后自动切回。')
p('切换是更新服务的内部行为，对回测引擎完全透明——引擎只读本地 Parquet。')

doc.add_heading('4.5 更新策略与耗时预估', level=2)
table(
    ['阶段', '内容', '预估耗时'],
    [
        ['首次全量', '5年日线 + 5年5分钟线，mootdx 为主', '约 1~2 小时（挂机一次）'],
        ['每日增量', '收盘后定时任务：当日日线 + 当日48根5分钟线 + 复权因子 + 交易日历', '几分钟至1小时内'],
        ['失败处理', '失败股票进重试队列，次日按交易日历缺口检测补拉（断点续传，不重复全量）', '—'],
    ],
    widths=[3, 9, 4],
)
p('数据质量校验（落盘前强制执行）：')
bullet('K 线数量 vs 交易日历比对（缺口检测）。')
bullet('异常价格检测：越界涨跌停、零成交、负价。')
bullet('复权因子连续性检查。')
bullet('校验通过才写入 Parquet，避免脏数据入湖。')

doc.add_heading('4.6 本地开发代理（Clash Verge）问题处理', level=2)
p('问题本质：Clash 开系统代理后，Python requests/httpx 默认读取 http_proxy 环境变量，导致国内数据源流量绕道海外节点——变慢、触发数据源风控、节点 IP 被封。baostock 走自研 TCP 协议一般不受影响，akshare/东财/腾讯走 HTTP 全部中招。')
p('解决方案（建议 A + B 同时做）：')
table(
    ['方案', '做法', '适用'],
    [
        ['A. 代码层强制直连（推荐，根治）', '所有数据源客户端统一使用 session.trust_env = False，无视系统代理与环境变量', '写进代码，保证程序行为在任何环境一致，避免"本地能跑、上云行为不一致"'],
        ['B. Clash 规则分流', 'Clash Verge 配置国内数据域名 DIRECT 规则（eastmoney.com、baostock.com、gtimg.cn、通达信服务器IP段）', '本机其他软件仍需代理时'],
        ['C. NO_PROXY 环境变量白名单', 'NO_PROXY=eastmoney.com,baostock.com,...', '作为 A 的兜底'],
    ],
    widths=[4.5, 7, 4.5],
)
p('云端部署无此问题（阿里云国内节点、无代理环境），但方案 A 必须写进代码以保证环境一致性。')

doc.add_page_break()

# ================= 5. 回测引擎设计 =================
doc.add_heading('5. 回测引擎设计', level=1)
doc.add_heading('5.1 事件驱动 + 向量化混合架构', level=2)
p('性能与灵活性的平衡点：')
bullet('信号层向量化：指标计算（MA/MACD/布林/ATR 等）用 polars 批量算完，产出"信号序列"。')
bullet('撮合层事件驱动：逐 bar 消费信号，处理订单提交 → 撮合 → 账户更新，精确模拟 T+1、涨跌停、手续费、滑点。')
code('DataFeed → SignalCalc(向量化) → Broker(事件驱动撮合) → Portfolio → TradeLog')

doc.add_heading('5.2 订单与分层持仓模型（支撑加减仓 / 做 T）', level=2)
p('核心设计——分层持仓账户：')
bullet('每笔开仓生成独立 Position（成本价、数量、开仓时间、可卖日期、归属标签）。')
bullet('账户 = Σ Positions，天然支持金字塔加仓（摊高成本）与分批止盈（只平部分仓位）。')
bullet('做 T 实现：底仓 Position 当日卖出 → 同日新买入生成新 Position → "卖旧买新"式做 T，每笔 T 的盈亏独立核算。')
p('A股真实约束撮合规则：')
table(
    ['规则', '实现'],
    [
        ['T+1', '当日买入的 Position 标记 sellable_date = 下一交易日'],
        ['涨跌停', '涨停价买单不成交、跌停价卖单不成交（防回测虚高收益）'],
        ['停牌/ST', '停牌期间无行情不撮合；ST 股可选过滤'],
        ['滑点', '固定比例 或 按成交量占比的冲击成本模型，可配置'],
        ['手续费', '佣金（最低5元）+ 印花税（卖出千1）+ 过户费，参数化'],
    ],
    widths=[3, 13],
)

doc.add_heading('5.3 风控模块（引擎内嵌，参数独立配置）', level=2)
bullet('个股仓位上限 / 总仓位上限。')
bullet('止损止盈：固定比例 / ATR 动态 / 移动止损（trailing stop）。')
bullet('动态阈值：盘中振幅阈值按近期 ATR 自适应（而非固定阈值）。')
bullet('最大回撤熔断：策略级，触发后停止开新仓。')
bullet('日内交易次数限制（防做 T 过度交易）。')
p('风控参数与策略参数分离，作为独立 risk_config 传入，便于 AI 寻优时单独扫描。')

doc.add_heading('5.4 统计与交易过程输出（AI 的原料）', level=2)
table(
    ['层级', '内容'],
    [
        ['交易明细 trade_log', '每笔买卖：时间/价格/数量/手续费/信号理由/所属建仓平仓组/交易类型（建仓|加仓|减仓|做T|止损）'],
        ['持仓快照', '每日收盘：持仓明细、市值、可用资金'],
        ['资金曲线', '每日净值、回撤、仓位利用率'],
        ['绩效报告', '年化收益、最大回撤、夏普/索提诺/卡玛、胜率、盈亏比、平均持仓天数、做T单独贡献、加减仓贡献分解、月度收益热力图'],
    ],
    widths=[4, 12],
)
p('做 T 贡献分解是核心指标：平仓按持仓时长 < 1 个交易日归类为 T 交易，单独统计胜率与盈亏，用于检验做 T 策略有效性；加减仓贡献分解同理。', bold=True)

doc.add_page_break()

# ================= 6. AI 辅助调优 =================
doc.add_heading('6. AI 辅助调优（多 LLM 配置）', level=1)
doc.add_heading('6.1 多 LLM Provider 架构', level=2)
p('系统不绑定单一 AI 服务商，采用 OpenAI 兼容协议的统一 Provider 抽象层，支持多配置并存、按需切换、自动降级：')
table(
    ['Provider', '协议', '用途建议'],
    [
        ['硅基流动 (SiliconFlow)', 'OpenAI 兼容', '主力：GLM/Qwen/DeepSeek 系列，有免费额度'],
        ['DeepSeek 官方', 'OpenAI 兼容', '便宜且强，适合报告深度分析'],
        ['智谱 GLM', 'OpenAI 兼容', '备选主力'],
        ['OpenAI / Claude', '各自 SDK 或兼容层', '可选，海外服务需网络条件'],
        ['Ollama 本地', 'OpenAI 兼容', '零成本、无速率限制，适合批量轻量任务（注意云服务器需有 GPU 或用小模型）'],
    ],
    widths=[4, 4, 8],
)
p('配置设计（详见第 9 章安全机制——api_key 一律走环境变量，配置文件只存 base_url 与模型名）：')
code(
'# llm.yaml（入库，不含密钥）\n'
'profiles:\n'
'  main:            # 默认分析模型\n'
'    provider: openai_compatible\n'
'    base_url: https://api.siliconflow.cn/v1\n'
'    model: deepseek-ai/DeepSeek-V3\n'
'    api_key_env: SILICONFLOW_API_KEY   # 密钥从环境变量读\n'
'  cheap:           # 轻量任务（分类/摘要）\n'
'    provider: openai_compatible\n'
'    base_url: http://localhost:11434/v1\n'
'    model: qwen2.5:7b\n'
'    api_key_env: OLLAMA_API_KEY\n'
'default: main\n'
'fallback_chain: [main, cheap]   # 主模型失败自动降级')
p('调度规则：')
numbered('不同任务用不同 profile：报告深度分析用 main，交易明细采样分类用 cheap，控制成本。')
numbered('失败降级：主模型超时/限流 → 沿 fallback_chain 自动切换。')
numbered('前端 AI 面板支持手动切换当前 profile，方便对比不同模型的分析质量。')
numbered('所有 LLM 调用记录 token 消耗与耗时，面板可见用量统计。')

doc.add_heading('6.2 AI 调优三层递进', level=2)
doc.add_heading('第一层：LLM 回测报告解读（最快见效）', level=3)
bullet('回测完成后，将结构化报告（绩效指标 + 交易明细摘要 + 资金曲线特征点）按统一 Prompt 模板喂给 LLM。')
bullet('输出：策略弱点诊断（如"止损过紧导致高换手低收益""做T胜率仅41%且贡献为负"）、参数敏感性判断（结合 Optuna 参数重要性）、具体优化建议（改哪个参数、往哪个方向、为什么）。')
bullet('交易明细过长时采样：盈利最多 / 亏损最多 / 随机抽样各 N 笔。')
doc.add_heading('第二层：AI + Optuna 参数寻优闭环', level=3)
code(
'回测 → LLM分析报告 → LLM给出参数搜索空间建议(范围缩放)\n'
'  ↑                                    ↓\n'
'  └──── 最优参数自动重跑验证 ←── Optuna贝叶斯寻优(带剪枝)')
bullet('LLM 不直接猜参数值，而是缩小搜索空间（如止损从 2%~20% 缩到 5%~9%），数值寻优交给 Optuna——各干擅长的事。')
bullet('Optuna MedianPruner 中途剪掉差的试验，节省算力。')
doc.add_heading('第三层：Agent 自主迭代（可选，后期）', level=3)
bullet('给 LLM 接 function calling 工具集：run_backtest(params) / get_report(id) / run_optimize(space)。')
bullet('Agent 自主循环"假设 → 回测 → 分析 → 调整"，设最大迭代次数与终止条件，防止过拟合循环。')
bullet('工具接口即第一/二层已有 API，改造成本低。')

doc.add_heading('6.3 防过拟合机制（必须做）', level=2)
bullet('样本内/外划分：寻优只用前 70% 数据，后 30% 样本外验证，报告两段对比。')
bullet('参数平原检查：最优参数 ±10% 微扰重跑，收益骤降即过拟合信号。')
bullet('AI 面板强制展示样本外表现，防止被漂亮曲线误导。')

doc.add_page_break()

# ================= 7. 前后端设计 =================
doc.add_heading('7. 前后端设计', level=1)
doc.add_heading('7.1 后端 API（FastAPI）', level=2)
code(
'POST   /api/auth/login                # JWT 登录\n'
'POST   /api/backtests                 # 创建回测任务 → task_id\n'
'GET    /api/backtests/{id}/status     # 状态查询（另有 WS /ws/tasks/{id} 推进度）\n'
'GET    /api/backtests/{id}/report     # 绩效报告+交易明细+资金曲线\n'
'GET    /api/backtests/{id}/kline      # K线+买卖点标记数据\n'
'POST   /api/optimize                  # 发起 Optuna 寻优\n'
'GET    /api/optimize/{id}/trials      # 试验结果\n'
'POST   /api/ai/analyze                # 触发 AI 分析某回测报告\n'
'GET    /api/ai/profiles               # LLM 配置列表/切换\n'
'GET    /api/data/status               # 数据更新水位/健康状态')
p('流程：前端提交 → 进程池执行回测 → 内存/磁盘写进度 → WebSocket 推前端 → 完成拉取报告渲染。')

doc.add_heading('7.2 前端页面', level=2)
table(
    ['页面', '要点'],
    [
        ['策略管理', '策略参数表单（参数 schema 驱动动态渲染）；可选 Monaco 代码编辑器'],
        ['回测控制台', '股票池、时间区间、周期（日线/5分钟）、初始资金、滑点手续费配置'],
        ['结果查看', 'KLineCharts 渲染 K 线+买卖点标记（开仓▲平仓▼，颜色区分建仓/加仓/做T）；资金曲线、回撤曲线'],
        ['统计报告', '指标卡片 + 月度热力图 + 交易明细表格（可筛选只看做T/加仓笔）'],
        ['寻优面板', 'Optuna 参数重要性图、平行坐标图、最优参数组合'],
        ['AI 面板', '对话式分析、优化建议列表（可一键"应用建议重跑"）、LLM 配置切换、token 用量统计'],
        ['数据管理', '数据更新水位、源健康状态、手动补拉入口'],
    ],
    widths=[3, 13],
)

doc.add_page_break()

# ================= 8. 部署方案 =================
doc.add_heading('8. 部署方案（阿里云 + Docker + 外网访问）', level=1)
doc.add_heading('8.1 服务器选型', level=2)
table(
    ['配置', '建议', '说明'],
    [
        ['机型', '阿里云轻量应用服务器 2C4G', '个人回测够用；活动价约 100~200 元/年'],
        ['系统盘', '默认 40~60GB', '系统+代码+镜像'],
        ['数据盘', 'ESSD entry 100GB', '5年数据约 45GB + 余量（Parquet+SQLite+寻优DB）'],
        ['带宽', '峰值 3~5Mbps', 'K线单次请求几百KB，几个用户足够'],
        ['地域', '国内节点（如杭州/上海）', '低延迟、访问国内数据源无代理问题'],
    ],
    widths=[3, 5, 8],
)

doc.add_heading('8.2 Docker 部署', level=2)
p('个人规模不需要 K8s，单个 docker-compose.yml 解决：')
code(
'docker-compose.yml\n'
'  services:\n'
'    app:        # FastAPI + 前端静态文件(单容器)\n'
'      build: .\n'
'      volumes:\n'
'        - ./data:/app/data          # Parquet数据湖(挂载,不入镜像)\n'
'        - ./config:/app/config      # 配置(含llm.yaml等,不入库)\n'
'        - ./.env:/app/.env:ro       # 密钥文件(不入库)\n'
'      environment:\n'
'        - SILICONFLOW_API_KEY=${SILICONFLOW_API_KEY}\n'
'    caddy:      # 反代+自动HTTPS\n'
'      ports: ["80:80", "443:443"]')
bullet('镜像分层：python-base → deps(requirements) → app(代码)，代码改动只重建最上层。')
bullet('本地开发与云端使用同一镜像，杜绝环境不一致。')
bullet('数据、配置、密钥全部通过 volume/env 注入，容器本身无状态、可随时重建。')

doc.add_heading('8.3 外网访问（异地朋友）', level=2)
numbered('域名：购买域名（约30元/年）解析到服务器公网 IP。')
numbered('HTTPS：Caddy 自动申请并续期 Let\'s Encrypt 证书，零配置。')
numbered('认证：FastAPI JWT 登录，朋友各开账号；或简化为统一访问码（个人系统不过度设计）。')
numbered('安全：仅开放 80/443 端口；SSH 改密钥登录并改端口；数据库不暴露（SQLite/DuckDB 本为本地文件）。')

doc.add_heading('8.4 定时任务', level=2)
bullet('每日收盘后（如 16:00）增量更新数据（mootdx + baostock，自动降级）。')
bullet('更新完成推送通知（可选：微信推送/邮件）。')
bullet('容器内用 APScheduler 或宿主机 cron 均可。')

doc.add_page_break()

# ================= 9. 安全与配置管理 =================
doc.add_heading('9. 安全与配置管理（GitHub 开源安全）', level=1)
p('项目要上传 GitHub，核心原则：代码入库，一切敏感信息（API Key、密码、数据文件）一律不入库。', bold=True)

doc.add_heading('9.1 敏感信息分层', level=2)
table(
    ['类别', '内容', '存放方式'],
    [
        ['密钥', 'LLM API Key、JWT Secret、通知 webhook', '.env 文件（.gitignore 排除）→ Docker env 注入；本地开发用 .env，服务器上同样'],
        ['业务配置', '数据源优先级、LLM profile（不含key）、回测默认参数', 'config/*.yaml；config.example.yaml 入库做模板，真实 config/ 目录 .gitignore'],
        ['行情数据', 'Parquet 数据湖、SQLite 元数据', '/data 目录整体 .gitignore（45GB 也不适合入库）；提供 data_builder 脚本让使用者自行拉取'],
        ['代码', '引擎、API、前端', '正常入库'],
    ],
    widths=[2.5, 6, 7.5],
)

doc.add_heading('9.2 .gitignore 关键条目', level=2)
code(
'.env\n'
'.env.*\n'
'!.env.example\n'
'config/\n'
'!config.example/\n'
'data/\n'
'*.db\n'
'*.sqlite3\n'
'*.parquet\n'
'logs/\n'
'__pycache__/\n'
'node_modules/\n'
'dist/')

doc.add_heading('9.3 防泄漏机制（纵深防御）', level=2)
numbered('.env.example 模板入库：列出全部所需环境变量名并留空注释，使用者复制为 .env 填写。')
numbered('pre-commit hook + gitleaks：每次提交前自动扫描密钥特征，命中即阻断提交。')
numbered('GitHub Actions 可选加 secret 扫描（gitleaks action）双保险。')
numbered('历史清洗：万一曾提交过密钥，用 git filter-repo 或 BFG 清洗历史，并立即作废该密钥（改密钥比删记录更重要）。')
numbered('LLM 配置的 api_key_env 机制：配置文件只写环境变量名，不出现密钥值（见 6.1 示例）。')
numbered('README 中明示"请勿提交 .env / data 目录"，并对协作者开启分支保护。')

doc.add_heading('9.4 环境变量清单（.env.example）', level=2)
code(
'# ---- LLM ----\n'
'SILICONFLOW_API_KEY=\n'
'DEEPSEEK_API_KEY=\n'
'ZHIPU_API_KEY=\n'
'OLLAMA_API_KEY=ollama\n'
'# ---- 应用 ----\n'
'JWT_SECRET=          # openssl rand -hex 32 生成\n'
'ADMIN_PASSWORD=\n'
'# ---- 通知(可选) ----\n'
'WXPUSH_TOKEN=')

doc.add_page_break()

# ================= 10. 性能优化 =================
doc.add_heading('10. 性能优化', level=1)
table(
    ['优化点', '手段', '预期收益', '优先级'],
    [
        ['指标计算', 'polars 替代 pandas（多线程）', '5~10x', 'P1（默认启用）'],
        ['数据读取', 'Parquet 按需 lazy-load + 进程内 LRU 缓存（策略池通常几十只票）', '分钟线场景 IO 大头', 'P1'],
        ['参数寻优', '进程池并行 trials + Optuna MedianPruner 剪枝', '近线性加速', 'P1'],
        ['大股票池回测', '按股票分片，多进程各跑各的再合并账户（兼作粗粒度断点续测）', '接近线性', 'P2'],
        ['热数据缓存', '复权后日线、常用指标结果缓存（diskcache）', '重复回测秒级启动', 'P2'],
        ['撮合循环', 'Numba JIT 编译逐bar撮合核心', '10~50x（仅循环部分）', 'P3（分钟线大池仍慢时才上）'],
    ],
    widths=[3.5, 6.5, 4, 2],
)
p('原则：先 polars + Parquet 分片（不改架构即有大收益）；Numba 等到真实瓶颈出现再上，避免一开始过度优化。')

# ================= 11. 断点续测 =================
doc.add_heading('11. 断点续测策略', level=1)
p('结论：单次回测不做精确断点续跑；参数寻优必须支持（Optuna 原生自带，零成本）。', bold=True)
doc.add_heading('11.1 单次回测：不做，理由', level=2)
numbered('单次没那么慢：3~5年日线回测秒级~分钟级；几十只票的 5 分钟线回测约 1~5 分钟，重跑代价小。')
numbered('全市场回测的慢场景用"分片并行"覆盖：按股票分片多进程，单片失败只重跑该片——本身就是粗粒度断点续测，实现成本低。')
numbered('精确续跑代价过高：引擎完整状态 = 所有持仓快照 + 资金曲线 + 未完成订单流 + 指标中间态，全量序列化会污染引擎代码，且极易出现"续跑结果与一次跑完不一致"的隐蔽 bug——比中断重跑更伤。')
p('折中预案（被痛点逼到再做）：按自然年 checkpoint，每年收盘落盘一次，中断后从最近整年恢复。')
doc.add_heading('11.2 参数寻优：Optuna 原生断点续跑', level=2)
code(
'进程挂了 → 重启 → 加载同一 study.db → 自动从第 N+1 个 trial 继续\n'
'已完成的 trial 结果全在，不浪费一次计算')
p('500 次试验 × 每次几分钟 = 十几小时的寻优任务，中断续跑是刚需，Optuna 的 SQLite 存储天然满足。')

doc.add_page_break()

# ================= 12. 实施路线图 =================
doc.add_heading('12. 实施路线图', level=1)
table(
    ['阶段', '内容', '交付物'],
    [
        ['P1 基础', '数据更新服务（mootdx+baostock 多源切换、增量、校验）+ Parquet 存储 + 日线回测引擎（订单/账户/风控）', '命令行跑通一个完整策略'],
        ['P2 平台', 'FastAPI + 进程池任务化 + React 回测控制台 + K线/报告展示 + JWT 认证', '完整前后端回测平台'],
        ['P3 核心', '5分钟线支持 + 做T/加减仓分层持仓 + 交易贡献分解统计', '满足核心业务需求'],
        ['P4 AI', '多 LLM Provider 配置 + LLM 报告分析 + Optuna 寻优 + AI 面板', 'AI 辅助闭环'],
        ['P5 完善', '阿里云 Docker 部署 + Caddy HTTPS + 定时更新 + gitleaks 安全加固 + 防过拟合工具', '上线给朋友用'],
        ['P6 进阶（可选）', 'Numba 加速 + Agent 自主迭代 + 样本外验证工具 + 文本事件研究（bge-m3 embedding，独立模块）', '完整形态'],
    ],
    widths=[2.5, 9, 4.5],
)

# ================= 13. 关键设计决策汇总 =================
doc.add_heading('13. 关键设计决策汇总', level=1)
table(
    ['#', '决策', '理由'],
    [
        ['1', '自研事件驱动引擎 + 向量化信号', '做T/分层加减仓/T+1/涨跌停精确模拟是第三方框架做不到位的部分'],
        ['2', 'Parquet + DuckDB 而非关系库存行情', '3.2亿行5分钟线的最优解，零运维'],
        ['3', '后复权 + 存储复权因子', '保证回测结果可复现'],
        ['4', '风控参数与策略参数分离', '便于独立扫描寻优'],
        ['5', 'mootdx 主源 + baostock 主备分工 + 自动降级', '速度与稳定兼顾，接口失效不断服务'],
        ['6', 'AI 定位"分析+缩小搜索空间"，数值寻优交给 Optuna', '各展所长，结果可控'],
        ['7', '多 LLM Provider 抽象（OpenAI 兼容协议）', '不绑定单一服务商，可切换可降级可控成本'],
        ['8', 'SQLite + 进程池，不用 PG/Celery/Redis', '个人规模够用，复杂度降一档'],
        ['9', '单次回测不做精确断点续跑，寻优靠 Optuna', '投入产出比 + 防隐蔽bug'],
        ['10', '密钥一律环境变量，数据目录不入库，gitleaks 扫描', 'GitHub 开源安全底线'],
        ['11', '数据源客户端 trust_env=False', '根治 Clash 代理导致的本地/云端行为不一致'],
        ['12', 'Docker compose 单文件部署 + Caddy', '环境一致、HTTPS 零配置、迁移方便'],
    ],
    widths=[1, 6.5, 8.5],
)

doc.save(r'd:\Sanchez\AI\TraeProjects\quan_quant\A股个人量化回测系统-完整实现方案.docx')
print('DOCX generated OK')
